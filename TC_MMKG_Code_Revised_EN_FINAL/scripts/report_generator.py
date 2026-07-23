# -*- coding: utf-8 -*-
from __future__ import annotations

from pathlib import Path
from typing import Iterable, Optional
from docx import Document
from docx.shared import Inches


def generate_assessment_report(
    output_path: str,
    damage_report: str = "",
    chart_path: Optional[str] = None,
    reasoning_chain: Optional[Iterable[str]] = None,
    verification_result: Optional[str] = None,
    interaction_records: Optional[Iterable[str]] = None,
) -> str:
    doc = Document()
    doc.add_heading('Truck-Crane Telescopic-Boom Remaining Fatigue Life Assessment Report', 0)
    if damage_report:
        doc.add_heading('1. Fatigue Damage and Remaining Life', level=1)
        for line in str(damage_report).splitlines():
            doc.add_paragraph(line)
    if chart_path and Path(chart_path).exists():
        doc.add_picture(str(chart_path), width=Inches(6.2))
    if reasoning_chain:
        doc.add_heading('2. Analogical Reasoning Chain', level=1)
        for i, step in enumerate(reasoning_chain, 1):
            doc.add_paragraph(f'{i}. {step}')
    if verification_result:
        doc.add_heading('3. Three-Layer Plausibility Validation', level=1)
        doc.add_paragraph(str(verification_result))
    if interaction_records:
        doc.add_heading('4. Interaction Records', level=1)
        for item in interaction_records:
            doc.add_paragraph(str(item))
    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    return output_path
